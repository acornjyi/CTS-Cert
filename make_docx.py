"""
DataON CTS 2026-2028 신청 작업 계획서 Word 문서 생성 스크립트
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"D:\업무\01 연구데이터센터\2026\01 CTS\DataON_CTS_2026-2028_신청_작업계획.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_header_row(table, headers, bg="2E4A7A", fg="FFFFFF"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(fg)
        run.font.size = Pt(9)
        set_cell_borders(cell)

def fill_row(table, row_idx, values, bg=None, bold_col=None):
    row = table.rows[row_idx]
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        p = cell.paragraphs[0]
        run = p.runs[0] if p.runs else p.add_run(str(val))
        if not p.runs:
            cell.text = ""
            run = p.add_run(str(val))
        run.font.size = Pt(9)
        if bold_col is not None and i == bold_col:
            run.bold = True
        if bg:
            set_cell_bg(cell, bg)
        set_cell_borders(cell)

def add_section_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(9)
    return p

# ── document setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins (2.5 cm all sides)
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10)

# Heading styles
for level, size, color_hex in [(1, 14, '2E4A7A'), (2, 12, '2E4A7A'), (3, 11, '1F3461')]:
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Malgun Gothic'
    h.font.size = Pt(size)
    h.font.color.rgb = RGBColor.from_string(color_hex)
    h.font.bold = True

# ── TITLE PAGE ────────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('DataON CoreTrustSeal 2026–2028\n신청 작업 계획서')
run.font.name = 'Malgun Gothic'
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string('2E4A7A')

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('General Repository 유형  ·  작성일: 2026-05-12  ·  작성: DataOn Repository Management Team')
run.font.name = 'Malgun Gothic'
run.font.size = Pt(11)
run.font.color.rgb = RGBColor.from_string('595959')

doc.add_paragraph()
doc.add_paragraph()

# ── SECTION 0: 신청 유형 및 현황 요약 ───────────────────────────────────────────

add_section_heading(doc, '0. 신청 유형 및 현황 요약', 1)

add_paragraph(doc, '신청 유형: General Repository (다학제 국가 연구데이터 플랫폼)', bold=True)
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('DataON(https://dataon.kisti.re.kr)은 KISTI가 운영하는 국가 연구데이터 공개 플랫폼으로, 자연과학·공학·사회과학 등 다학제 분야 연구데이터를 수집·보존·제공합니다. CoreTrustSeal(CTS) Requirements 2026–2028 기준으로 17개 요건(R0–R16)에 대한 갱신 신청을 준비합니다.')
run.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(doc, '현재 준수 수준 요약', bold=True)

# Compliance summary table
t = doc.add_table(rows=4, cols=3)
t.style = 'Table Grid'

add_header_row(t, ['준수 수준', '해당 요건 수', '요건 번호'])
data = [
    ('✅ 준수 완료 (레벨 3–4)', '8', 'R0, R2, R3, R8, R10, R12, R13, R14'),
    ('🟡 부분 준수 — 보완 필요 (레벨 2)', '7', 'R1, R4, R5, R6, R7, R9, R11'),
    ('❌ 미준수 — 즉시 조치 필요 (레벨 0–1)', '2', 'R15, R16'),
]
for i, (a, b, c) in enumerate(data, 1):
    fill_row(t, i, [a, b, c])

doc.add_paragraph()

# ── SECTION 1: 준비해야 할 문서 ──────────────────────────────────────────────────

add_section_heading(doc, '1. 준비해야 할 문서', 1)

docs_list = [
    ('1.1 기관·서비스 설명 문서 [Critical]', [
        'DataON 미션·비전 선언문 (Scope of collections, Mission statement)',
        '조직도 및 거버넌스 구조 설명 (KISTI 내 DataON 위치 포함)',
        '서비스 통계 보고서: 사용자 수, 데이터셋 수, 다운로드 수 (최근 3년)',
        '재정 지속성 증빙: KISTI 예산 배정 내역 또는 중기재정계획',
    ]),
    ('1.2 정책 문서 [Critical]', [
        '데이터 수집 정책 (Collection Policy) — 수집 범위·기준·절차 명시',
        '데이터 접근 및 재이용 정책 (Access & Reuse Policy) — 라이선스 체계 포함',
        '장기보존 정책 (Preservation Policy) — 보존 기간·포맷 마이그레이션·폐기 기준',
        '데이터 품질 보증 정책 (Data Quality Policy) — 검수 절차·메타데이터 품질 기준',
    ]),
    ('1.3 기술 문서 [High]', [
        '시스템 아키텍처 다이어그램 (서버·스토리지·네트워크 구성)',
        '백업·복구 절차서 (백업 주기, RTO/RPO, 복구 테스트 결과)',
        '접근 통제 정책 (인증·권한 관리, 데이터 접근 로그)',
        '보안 점검 결과 보고서 (취약점 스캔, ISMS 또는 동등 인증)',
    ]),
    ('1.4 메타데이터 및 표준 문서 [High]', [
        '메타데이터 스키마 문서 (37개 필드 정의, 필수/선택 구분)',
        '사용 표준 목록: Dublin Core, DataCite, DDI, ISO 19115 등',
        'PID 시스템 설명: DOI 발급 절차 및 DOI 부여 정책',
    ]),
    ('1.5 데이터 관리 절차서 [High]', [
        '데이터 수집·등록 워크플로우 (제출→검수→승인→공개)',
        '버전 관리 정책 (데이터 수정·업데이트 이력 추적)',
        '데이터 폐기 절차 (철회 요청 처리, Tombstone 페이지 유지)',
    ]),
    ('1.6 이용자 지원 문서 [Medium]', [
        '이용자 가이드 및 FAQ',
        '데이터 기탁 매뉴얼 (연구자 대상)',
        '교육 프로그램 커리큘럼 (온라인 강좌, 워크숍)',
    ]),
    ('1.7 규정 준수 문서 [Medium]', [
        '개인정보 처리방침 (PIPA 준수)',
        '저작권 및 라이선스 정책 (CC 라이선스 적용 기준)',
        '윤리 검토 절차 (민감 데이터 처리 기준)',
    ]),
    ('1.8 성과 및 영향 문서 [Medium]', [
        '연간 서비스 보고서',
        '사용자 만족도 조사 결과',
        '데이터 인용 현황 (DataCite 인용 통계 등)',
    ]),
]

for title_bold, bullets in docs_list:
    add_paragraph(doc, title_bold, bold=True, size=10)
    for b in bullets:
        add_bullet(doc, b)
    doc.add_paragraph()

# ── SECTION 2: 신청서 내 보완 필요 항목 ─────────────────────────────────────────

add_section_heading(doc, '2. 신청서 내 보완 필요 항목', 1)

add_paragraph(doc, '아래 항목은 기존 신청서 초안에 [INSERT] 또는 [TO BE CONFIRMED] 표시가 있거나 구체적 근거가 부족한 부분입니다.', size=9)
doc.add_paragraph()

t2 = doc.add_table(rows=22, cols=4)
t2.style = 'Table Grid'
add_header_row(t2, ['요건', '항목', '현황', '필요 조치'])

items2 = [
    ('R0', '미션 선언문', '[INSERT]', '공식 미션 문구 확정 및 웹사이트 게시'),
    ('R1', '재정 지속성 증빙', '[TO BE CONFIRMED]', 'KISTI 중기재정계획 또는 예산 배정 공문 첨부'),
    ('R1', '직원 FTE 수', '[INSERT]', '현재 인력 현황표 작성 (정규직/계약직 구분)'),
    ('R2', '수집 범위 기준', '부분 기술', '분야별·형식별 수집 기준 문서화'),
    ('R3', '라이선스 정책', '부분 기술', 'CC 라이선스 선택 기준 및 의무 CC-BY 적용 범위 명시'),
    ('R4', '무결성 검증 절차', '[INSERT]', '체크섬(MD5/SHA-256) 생성·검증 자동화 여부 기술'),
    ('R5', '메타데이터 품질 절차', '부분 기술', '검수 체크리스트 및 담당자 역할 명시'),
    ('R6', '워크플로우 문서', '[INSERT]', '제출→검수→승인→공개 단계별 SLA 정의'),
    ('R7', '접근 통제', '[TO BE CONFIRMED]', '인증 방식, 권한 등급, 로그 보존 기간 명시'),
    ('R7', '개인정보 처리방침', '[INSERT]', 'PIPA 준수 방침 URL 및 최종 개정일 기재'),
    ('R8', '평가자 교육', '[INSERT]', '데이터 검수자 교육 이력 및 자격 요건 기술'),
    ('R9', '도구·소프트웨어 목록', '[TO BE CONFIRMED]', '사용 중인 리포지토리 소프트웨어, 버전, 라이선스 목록화'),
    ('R10', '보존 포맷 정책', '부분 기술', '권장 포맷 목록 및 마이그레이션 트리거 조건 문서화'),
    ('R11', '백업 절차', '[INSERT]', 'RTO/RPO 수치, 백업 주기, 오프사이트 복사 여부 명시'),
    ('R12', 'DOI 발급 정책', '부분 기술', 'DOI 부여 기준(데이터셋 크기·형식·공개 여부) 명시'),
    ('R13', '데이터 인용 지원', '기술됨', '인용 형식 예시 추가'),
    ('R14', '이용자 지원 채널', '기술됨', '응답 SLA 명시 (예: 영업일 기준 3일 이내)'),
    ('R15', 'OAIS 준수 선언', '❌ 없음', 'OAIS 기능 매핑 문서 작성 (6개 기능 엔티티)'),
    ('R15', 'Designated Community', '❌ 없음', 'DC 정의 문서 작성 (대상 이용자 집단, 지식 기반)'),
    ('R16', '지속적 개선 체계', '❌ 없음', '연간 자체점검 계획 및 개선 이력 관리 체계 수립'),
    ('R16', '이해관계자 피드백', '[INSERT]', '사용자 만족도 조사 결과 반영 절차 기술'),
]

for i, row_data in enumerate(items2, 1):
    req, item, status, action = row_data
    row = t2.rows[i]
    for j, val in enumerate([req, item, status, action]):
        cell = row.cells[j]
        cell.text = val
        p = cell.paragraphs[0]
        if p.runs:
            run = p.runs[0]
        else:
            run = p.add_run(val)
        run.font.size = Pt(9)
        if j == 2 and '❌' in val:
            set_cell_bg(cell, 'FFD7D7')
        elif j == 2 and '[INSERT]' in val:
            set_cell_bg(cell, 'FFF2CC')
        elif j == 2 and '[TO BE CONFIRMED]' in val:
            set_cell_bg(cell, 'FFF2CC')
        set_cell_borders(cell)

doc.add_paragraph()

# ── SECTION 3: Dead URL 교체 작업 ────────────────────────────────────────────────

add_section_heading(doc, '3. Dead URL 교체 작업', 1)

add_paragraph(doc, '2026-05-12 URL 검사 결과(총 96개) 중 교체가 필요한 DataON 내부 URL 및 외부 URL 목록입니다.', size=9)
doc.add_paragraph()

add_paragraph(doc, '3-1. DataON 내부 URL 교체 목록', bold=True, size=10)

t3a = doc.add_table(rows=10, cols=4)
t3a.style = 'Table Grid'
add_header_row(t3a, ['각주 번호', '구 URL (404)', '새 URL', '비고'])

dead_dataon = [
    ('1, 21, 62', 'intro/intro01.do?mm=MENU_00102&sm=MENU00144', 'https://dataon.kisti.re.kr/portal/intro/outline.do', 'DataON 소개'),
    ('4', 'intro/intro03.do?mm=MENU_00102&sm=MENU00146', 'https://dataon.kisti.re.kr/portal/intro/history.do', '개발 연혁'),
    ('5', 'intro/intro02.do?mode=1&mm=MENU_00102&sm=MENU00145', 'https://dataon.kisti.re.kr/portal/intro/outline.do', '서비스 소개'),
    ('6', 'intro/intro07.do?mm=MENU_00102&sm=MENU00200', 'https://dataon.kisti.re.kr/portal/intro/assnOrg.do', '협력 기관'),
    ('33, 44', 'intro/intro08.do', 'https://dataon.kisti.re.kr/portal/intro/mgmtPolicy.do', '조직도/거버넌스'),
    ('48', 'kacademy.kisti.re.kr/online-edu/free/620b7cf8...', 'https://kacademy.kisti.re.kr/', 'KISTI 교육 홈'),
    ('66', 'intro/intro01.do', 'https://dataon.kisti.re.kr/portal/intro/outline.do', '데이터 수집 범위'),
    ('67, 94', 'intro/intro09.do?mm=MENU_00102&sm=MENU00307', 'https://dataon.kisti.re.kr/portal/intro/catalog_summary.do', '메타데이터 스키마'),
    ('51', 'rd-alliance.org/...rda-technical-advisory-board.html', 'https://www.rd-alliance.org/about-rda/rda-leadership/', 'RDA TAB'),
]

for i, row_data in enumerate(dead_dataon, 1):
    row = t3a.rows[i]
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        cell.text = val
        p = cell.paragraphs[0]
        run = p.runs[0] if p.runs else p.add_run(val)
        run.font.size = Pt(8)
        set_cell_borders(cell)

doc.add_paragraph()
add_paragraph(doc, '3-2. 외부 URL 교체 필요 목록', bold=True, size=10)

t3b = doc.add_table(rows=3, cols=4)
t3b.style = 'Table Grid'
add_header_row(t3b, ['각주 번호', '구 URL', '상태', '조치'])

ext_dead = [
    ('93', 'tta.or.kr/tta/ttaSearchView.do?...TTAK.KO-10.0976', '410 Gone', 'TTA 표준 검색 페이지에서 TTAK.KO-10.0976 수동 검색 후 신규 URL로 교체'),
    ('50', 'scidatacon.org/IDW-2022/sessions/293/', '404', 'Wayback Machine 아카이브 URL 또는 발표자료 DOI로 대체'),
]

for i, row_data in enumerate(ext_dead, 1):
    row = t3b.rows[i]
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        cell.text = val
        p = cell.paragraphs[0]
        run = p.runs[0] if p.runs else p.add_run(val)
        run.font.size = Pt(8)
        set_cell_borders(cell)

doc.add_paragraph()

add_paragraph(doc, '* dmp.kisti.re.kr (각주 18, 70b, 89b): 검사 당일 서버 다운 상태였으므로, 재점검 후 신규 DMP 서비스 URL로 교체 필요.', size=8)

doc.add_paragraph()

# ── SECTION 4: General Repository 특이사항 ──────────────────────────────────────

add_section_heading(doc, '4. General Repository 유형 특이사항', 1)

add_paragraph(doc, 'CoreTrustSeal은 General Repository에 대해 아래 사항을 특별히 심사합니다.', size=9)
doc.add_paragraph()

t4 = doc.add_table(rows=7, cols=3)
t4.style = 'Table Grid'
add_header_row(t4, ['심사 항목', '현황', '보완 방향'])

gen_repo = [
    ('다학제 데이터 처리 역량', '자연과학·공학 위주', '사회과학·인문학 데이터 수용 정책 명시'),
    ('Designated Community 다양성', '미정의', '연구자·학생·정책입안자 등 DC 세분화 문서 작성'),
    ('큐레이션 품질 기준', '부분 문서화', '분야별 메타데이터 품질 기준 차등 적용 방안 수립'),
    ('데이터 형식 폭넓은 수용', 'CSV, Excel, HDF5 등 지원', '비정형·대용량 포맷(netCDF, Parquet) 지원 계획 추가'),
    ('OAIS 기능 매핑', '없음', '제출 정보 패키지(SIP)→보존 정보 패키지(AIP)→배포 정보 패키지(DIP) 흐름 문서화'),
    ('접근 제한 데이터 처리', '공개 위주', '비공개·접근 제한 데이터 처리 절차 명시 (연구 목적 별도 신청 등)'),
]

for i, row_data in enumerate(gen_repo, 1):
    fill_row(t4, i, row_data)

doc.add_paragraph()

# ── SECTION 5: 작업 일정 ───────────────────────────────────────────────────────

add_section_heading(doc, '5. 작업 일정', 1)

t5 = doc.add_table(rows=5, cols=4)
t5.style = 'Table Grid'
add_header_row(t5, ['단계', '기간', '주요 작업', '산출물'])

timeline = [
    ('1단계 — 현황 분석 및 Gap 파악', '2026-05 ~ 06', 'URL 교체, 신청서 초안 검토, [INSERT] 항목 목록화', '보완 항목 목록, 교체 URL 목록'),
    ('2단계 — 정책 문서 정비', '2026-06 ~ 07', '수집 정책, 보존 정책, 접근 정책 초안 작성', '정책 문서 3종'),
    ('3단계 — 기술 문서 보완', '2026-07 ~ 08', '시스템 아키텍처 문서, 백업 절차서, OAIS 매핑 작성', '기술 문서 3종'),
    ('4단계 — 신청서 최종 작성 및 검토', '2026-08 ~ 09', '전 요건 응답 완성, 내부 검토, 영문 교정', '완성 신청서'),
]

for i, row_data in enumerate(timeline, 1):
    fill_row(t5, i, row_data)

doc.add_paragraph()

# ── SECTION 6: 역할 분담 ───────────────────────────────────────────────────────

add_section_heading(doc, '6. 역할 분담', 1)

t6 = doc.add_table(rows=8, cols=3)
t6.style = 'Table Grid'
add_header_row(t6, ['담당 영역', '담당자/팀', '비고'])

roles = [
    ('신청서 총괄 조율', 'Repository Manager', 'CTS 요건 이해, 전체 일정 관리'),
    ('정책 문서 작성', '정책·기획팀', '수집·보존·접근·품질 정책'),
    ('기술 문서 작성', '시스템 운영팀', '아키텍처, 백업, 보안'),
    ('메타데이터 문서', '데이터 큐레이션팀', '스키마, 품질 절차'),
    ('Dead URL 교체', 'Repository Manager + 웹팀', '각주 URL 일괄 교체'),
    ('영문 교정', '외부 전문가 또는 내부 영어 능통자', '신청서 영문 응답 교정'),
    ('최종 제출', 'Repository Manager', 'CTS 포털 업로드 및 제출'),
]

for i, row_data in enumerate(roles, 1):
    fill_row(t6, i, row_data)

doc.add_paragraph()

# ── SECTION 7: 최종 제출 전 체크리스트 ──────────────────────────────────────────

add_section_heading(doc, '7. 최종 제출 전 체크리스트', 1)

t7 = doc.add_table(rows=16, cols=3)
t7.style = 'Table Grid'
add_header_row(t7, ['항목', '완료 여부', '담당'])

checklist = [
    ('R0–R16 전 요건 응답 작성 완료', '☐', '총괄'),
    ('[INSERT]/[TO BE CONFIRMED] 항목 전부 해소', '☐', '총괄'),
    ('Dead URL 전부 교체 (각주 내 URL 96개 재검사)', '☐', 'Repository Manager'),
    ('수집 정책 문서 첨부', '☐', '정책팀'),
    ('보존 정책 문서 첨부', '☐', '정책팀'),
    ('접근·재이용 정책 문서 첨부', '☐', '정책팀'),
    ('데이터 품질 정책 문서 첨부', '☐', '큐레이션팀'),
    ('시스템 아키텍처 다이어그램 첨부', '☐', '운영팀'),
    ('백업·복구 절차서 첨부', '☐', '운영팀'),
    ('OAIS 기능 매핑 문서 첨부', '☐', 'Repository Manager'),
    ('Designated Community 정의 문서 첨부', '☐', 'Repository Manager'),
    ('재정 지속성 증빙 자료 첨부', '☐', '기획팀'),
    ('개인정보 처리방침 URL 확인', '☐', '운영팀'),
    ('신청서 영문 교정 완료', '☐', '영문 교정 담당'),
    ('CTS 포털 제출 완료', '☐', '총괄'),
]

for i, row_data in enumerate(checklist, 1):
    row = t7.rows[i]
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        cell.text = val
        p = cell.paragraphs[0]
        run = p.runs[0] if p.runs else p.add_run(val)
        run.font.size = Pt(9)
        if j == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_borders(cell)

doc.add_paragraph()

# ── FOOTER NOTE ──────────────────────────────────────────────────────────────

add_paragraph(doc, '본 계획서는 DataON Repository Management Team이 작성하였으며, CTS Requirements 2026–2028 기준을 따릅니다.', size=8)
add_paragraph(doc, '참고: CoreTrustSeal — https://www.coretrustseal.org/', size=8)

# ── SAVE ─────────────────────────────────────────────────────────────────────

doc.save(OUTPUT_PATH)
print(f"문서 저장 완료: {OUTPUT_PATH}")
